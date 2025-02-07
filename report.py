from docx import Document
from docx.shared import Inches
from config import REPORT_CONFIG
import os
import matplotlib.pyplot as plt
from io import BytesIO
import tempfile
from datetime import datetime

class ReportGenerator:
    def __init__(self):
        self.document = Document()
        # 使用临时目录
        self.temp_dir = tempfile.gettempdir()

    def add_heading(self, text, level=1):
        """添加标题"""
        self.document.add_heading(text, level=level)

    def add_paragraph(self, text):
        """添加段落"""
        self.document.add_paragraph(text)

    def add_chart(self, figure, caption=""):
        try:
            # 保存图表
            img_stream = BytesIO()
            figure.savefig(img_stream, format='png', 
                         bbox_inches='tight', 
                         dpi=300,
                         bbox_extra_artists=(figure.texts + figure.legends + figure.axes))
            img_stream.seek(0)
            
            # 添加图表到文档，设置合适的宽度
            width = Inches(5.0)  # 设置图片宽度为6英寸
            self.document.add_picture(img_stream, width=width)
            
            # 添加图表说明
            if caption:
                self.document.add_paragraph(caption)
            
            # 清理
            plt.close(figure)
            img_stream.close()
            
        except Exception as e:
            print(f"添加图表时出错: {str(e)}")
            raise

    def save_report(self, output_file='report.docx'):
        try:
            # 获取完整的目标路径
            target_path = os.path.abspath(output_file)
            
            # 检查文件是否已存在且被占用
            if os.path.exists(target_path):
                try:
                    # 尝试打开文件
                    with open(target_path, 'a'):
                        pass
                except PermissionError:
                    raise PermissionError(f"文件 '{target_path}' 已被其他程序打开，请关闭后重试")
            
            # 确保目标目录存在
            target_dir = os.path.dirname(target_path)
            if not os.path.exists(target_dir):
                try:
                    os.makedirs(target_dir)
                except PermissionError:
                    raise PermissionError(f"没有权限创建目录 '{target_dir}'，请检查权限设置")
                except Exception as e:
                    raise Exception(f"创建目录失败: {str(e)}")
            
            # 尝试保存文件
            try:
                self.document.save(target_path)
                print(f"报告已成功保存到: {target_path}")
            except PermissionError:
                raise PermissionError(f"没有权限保存到 '{target_path}'，请检查文件权限或尝试使用管理员权限运行")
            
        except Exception as e:
            error_msg = str(e)
            if "Permission denied" in error_msg:
                print("保存失败：权限不足")
                print("可能的原因：")
                print("1. 文件被其他程序（如 Word）打开")
                print("2. 没有文件夹的写入权限")
                print("3. 需要管理员权限")
                print("\n建议：")
                print("- 关闭已打开的 Word 文档")
                print("- 检查文件夹权限")
                print("- 尝试使用管理员权限运行程序")
            else:
                print(f"保存报告时出错: {error_msg}")
            raise

    # 设置 save 为 save_report 的别名
    save = save_report 